"""Docx extension."""
from .IngestorInterface import IngestorInterface
from typing import List
from .QuoteModel import QuoteModel
from docx import Document


class docxDocument(IngestorInterface):
    """docxDocument class."""

    extensions = ['docx']

    @classmethod
    def parse(cls, path: str) -> List[QuoteModel]:
        """Parse method with docx file."""
        if not cls.can_ingest(path):
            raise Exception('The format is unsupported.')
        try:
            ls_docx = []
            lines = Document(path)
            for line in lines.paragraphs:
                p = line.text.split(sep=' - ')
                p = [s.strip() for s in p]
                if p[0] != "":
                    ls_docx.append(QuoteModel(p[0], p[1]))
            return ls_docx
        except Exception as e:
            print(f'Error: {e}')
            return []
